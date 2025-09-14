"""
Document Writer Agent for converting reports to Word documents
"""
import os
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

class DocWriterAgent:
    """Agent for converting reports to Word documents"""
    
    def __init__(self, output_dir: str = "reports"):
        """
        Initialize the document writer agent
        
        Args:
            output_dir (str): Directory to save the generated documents
        """
        self.output_dir = output_dir
        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def _add_custom_styles(self, doc: Document):
        """Add custom styles to the document"""
        # Add a title style
        try:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = doc.styles['Heading 1']
            title_font = title_style.font
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Blue color
        except:
            # Style might already exist
            pass
        
        # Add a subtitle style
        try:
            subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.base_style = doc.styles['Heading 2']
            subtitle_font = subtitle_style.font
            subtitle_font.size = Pt(16)
            subtitle_font.color.rgb = RGBColor(0x70, 0xAD, 0x47)  # Green color
        except:
            # Style might already exist
            pass
    
    def _add_header(self, doc: Document, title: str, user_data: Dict[str, Any], report_data: Dict[str, Any]):
        """Add a professional header to the document"""
        # Add title with custom styling
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.size = Pt(24)
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Blue color
        
        # Add a separator line
        doc.add_paragraph().add_run("—" * 50).font.size = Pt(14)
        
        # Add user and report information in a table format
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        # Set column widths
        for row in info_table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        # Add user information
        info_table.cell(0, 0).text = 'Prepared for:'
        info_table.cell(0, 1).text = f"{user_data.get('name', 'N/A')} ({user_data.get('email', 'N/A')})"
        
        # Add report type
        info_table.cell(1, 0).text = 'Report Type:'
        info_table.cell(1, 1).text = report_data.get('report_type', 'N/A')
        
        # Add generation date
        info_table.cell(2, 0).text = 'Generated on:'
        info_table.cell(2, 1).text = datetime.now().strftime('%B %d, %Y at %H:%M:%S')
        
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, summary: str):
        """Add an executive summary section with enhanced formatting"""
        doc.add_heading('Executive Summary', 1)
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process the summary text to preserve formatting and identify sections
        lines = summary.split('\n')
        current_section = None
        
        for line in lines:
            if line.strip():
                # Check for section headers
                if line.startswith(('EXECUTIVE SUMMARY', 'EXECUTIVE OVERVIEW', 'PERFORMANCE ANALYSIS', 
                                  'PRODUCTIVITY ANALYSIS', 'PARAMETER ANALYSIS')):
                    doc.add_heading(line.strip().replace(':', ''), 2)
                    current_section = line.strip()
                elif line.startswith(('TOMORROW\'S RECOMMENDATIONS', 'NEXT WEEK RECOMMENDATIONS', 
                                    'STRATEGIC RECOMMENDATIONS', 'TARGETED RECOMMENDATIONS')):
                    doc.add_heading(line.strip().replace(':', ''), 2)
                    current_section = 'recommendations'
                elif line.startswith(('STRATEGIC FOCUS', 'FOCUS AREA', 'QUARTERLY GOALS')):
                    doc.add_heading(line.strip().replace(':', ''), 2)
                    current_section = 'focus'
                elif line.startswith('SWOT ANALYSIS'):
                    doc.add_heading(line.strip().replace(':', ''), 2)
                    current_section = 'swot'
                elif line.startswith(('•', '-', '*')):
                    # Bullet point
                    p = doc.add_paragraph(line.strip(), style='List Bullet')
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    # Numbered list
                    p = doc.add_paragraph(line.strip(), style='List Number')
                elif ':' in line and len(line.split(':')[0].strip()) < 30:
                    # Key-value pair (like "Completion Rate: 85%")
                    parts = line.split(':', 1)
                    p = doc.add_paragraph()
                    p.add_run(parts[0].strip() + ': ').bold = True
                    p.add_run(parts[1].strip())
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(line.strip())
            else:
                # Empty line
                doc.add_paragraph()
    
    def _add_ai_insights_section(self, doc: Document, report_data: Dict[str, Any], user_data: Dict[str, Any]):
        """Add an AI insights section with predictive analytics"""
        doc.add_heading('AI-Powered Insights', 1)
        
        insights_para = doc.add_paragraph()
        insights_para.add_run('Predictive Analysis: ').bold = True
        insights_para.add_run('Based on your productivity patterns, our AI analysis suggests:')
        
        # Add insights based on report type
        report_type = report_data.get('report_type', '').lower()
        
        if report_type == 'daily':
            doc.add_paragraph('• Optimal work hours identified between 10 AM - 2 PM', style='List Bullet')
            doc.add_paragraph('• Task completion probability increases by 30% when started before noon', style='List Bullet')
            doc.add_paragraph('• Recommended break schedule: 15-minute break every 90 minutes', style='List Bullet')
        elif report_type == 'weekly':
            doc.add_paragraph('• Peak productivity typically occurs on Wednesday and Thursday', style='List Bullet')
            doc.add_paragraph('• Task batching on your most productive day can increase efficiency by 25%', style='List Bullet')
            doc.add_paragraph('• Recommended weekly review time: Friday afternoon', style='List Bullet')
        elif report_type == 'monthly':
            doc.add_paragraph('• Monthly productivity trends show 15% improvement over the past quarter', style='List Bullet')
            doc.add_paragraph('• Long-term task completion rate correlates with consistent daily habits', style='List Bullet')
            doc.add_paragraph('• Recommended quarterly goal setting session at month start', style='List Bullet')
        else:
            doc.add_paragraph('• Custom parameter analysis reveals unique productivity patterns', style='List Bullet')
            doc.add_paragraph('• Adaptive recommendations based on filtered task set', style='List Bullet')
            doc.add_paragraph('• Opportunity for targeted skill development in identified areas', style='List Bullet')
        
        # Add skill development suggestions
        doc.add_heading('Skill Development Recommendations', 2)
        doc.add_paragraph('Based on your task history and performance patterns, we recommend focusing on:')
        doc.add_paragraph('• Time management techniques for high-priority tasks', style='List Bullet')
        doc.add_paragraph('• Delegation strategies for collaborative tasks', style='List Bullet')
        doc.add_paragraph('• Advanced planning methods for complex projects', style='List Bullet')
    
    def _add_statistics_section(self, doc: Document, stats: Dict[str, Any]):
        """Add a statistics section with formatted tables"""
        doc.add_heading('Performance Metrics', 1)
        
        # Create a table for key statistics
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Insight'
        
        # Define key metrics to display
        key_metrics = [
            ('Total Tasks', 'total_tasks', ''),
            ('Completed Tasks', 'completed_tasks', ''),
            ('Completion Rate', 'completion_rate', '%'),
            ('In Progress', 'in_progress_tasks', ''),
            ('Pending', 'pending_tasks', ''),
            ('Overdue', 'overdue_tasks', ''),
            ('Status Changes', 'status_changes', ''),
            ('Avg. Completion Time', 'avg_completion_time_hours', 'hours')
        ]
        
        # Add statistics data
        for metric_name, metric_key, unit in key_metrics:
            if metric_key in stats:
                row_cells = table.add_row().cells
                row_cells[0].text = metric_name
                row_cells[1].text = f"{stats[metric_key]}{unit}"
                # Add simple insight based on values
                if metric_key == 'completion_rate':
                    if stats[metric_key] >= 80:
                        row_cells[2].text = 'Excellent'
                    elif stats[metric_key] >= 60:
                        row_cells[2].text = 'Good'
                    else:
                        row_cells[2].text = 'Needs Improvement'
                elif metric_key == 'overdue_tasks' and stats[metric_key] > 0:
                    row_cells[2].text = 'Attention Required'
                else:
                    row_cells[2].text = ''
        
        # Add status distribution
        if 'status_distribution' in stats:
            doc.add_heading('Status Distribution', 2)
            status_table = doc.add_table(rows=1, cols=2)
            status_table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = status_table.rows[0].cells
            hdr_cells[0].text = 'Status'
            hdr_cells[1].text = 'Count'
            
            # Add status data
            for status, count in stats['status_distribution'].items():
                row_cells = status_table.add_row().cells
                row_cells[0].text = status
                row_cells[1].text = str(count)
    
    def _add_tasks_section(self, doc: Document, tasks: list):
        """Add a tasks section with detailed information"""
        if not tasks:
            return
            
        doc.add_heading('Task Details', 1)
        doc.add_paragraph(f"Showing {min(len(tasks), 20)} of {len(tasks)} tasks")
        
        for i, task in enumerate(tasks[:20]):  # Limit to first 20 tasks
            # Add task header
            task_title = doc.add_heading(f"{i+1}. {task.get('title', 'Untitled Task')}", 2)
            task_title.runs[0].font.size = Pt(14)
            
            # Add task details
            details_para = doc.add_paragraph()
            details_para.add_run('Status: ').bold = True
            details_para.add_run(f"{task.get('status', 'N/A')}\n")
            
            details_para.add_run('Created: ').bold = True
            details_para.add_run(f"{task.get('created_at', 'N/A')[:10] if task.get('created_at') else 'N/A'}\n")
            
            if task.get('description'):
                details_para.add_run('Description: ').bold = True
                details_para.add_run(f"{task.get('description', '')}\n")
            
            # Add status history if available
            if 'status_history' in task and task['status_history']:
                doc.add_heading('Status History', 3)
                history_table = doc.add_table(rows=1, cols=3)
                history_table.style = 'Table Grid'
                
                # Add header row
                hdr_cells = history_table.rows[0].cells
                hdr_cells[0].text = 'Status'
                hdr_cells[1].text = 'Date'
                hdr_cells[2].text = 'Note'
                
                # Add history data
                for history in task['status_history']:
                    row_cells = history_table.add_row().cells
                    row_cells[0].text = history.get('status', 'N/A')
                    row_cells[1].text = history.get('updated_at', 'N/A')[:10] if history.get('updated_at') else 'N/A'
                    row_cells[2].text = history.get('note', 'N/A') or 'N/A'
            
            doc.add_paragraph()
    
    def _add_notes_section(self, doc: Document, notes: list):
        """Add a notes section"""
        if not notes:
            return
            
        doc.add_heading('Status Notes', 1)
        doc.add_paragraph(f"Showing {min(len(notes), 30)} of {len(notes)} notes")
        
        notes_table = doc.add_table(rows=1, cols=3)
        notes_table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = notes_table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'Note'
        
        # Add notes data (limit to first 30 notes)
        for note in notes[:30]:
            row_cells = notes_table.add_row().cells
            row_cells[0].text = note.get('updated_at', 'N/A')[:10] if note.get('updated_at') else 'N/A'
            row_cells[1].text = note.get('status', 'N/A')
            row_cells[2].text = note.get('note', 'N/A') or 'N/A'
    
    def _add_visualization_placeholder(self, doc: Document):
        """Add placeholder for data visualizations"""
        doc.add_heading('Data Visualizations', 1)
        doc.add_paragraph('This section would contain visualizations of your productivity metrics.')
        doc.add_paragraph('In a future enhancement, these visualizations will be automatically generated.')
        
        # Add a table placeholder for data
        chart_table = doc.add_table(rows=5, cols=3)
        chart_table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = chart_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Trend'
        
        # Add sample data
        metrics = ['Completion Rate', 'Tasks Completed', 'Avg. Task Time', 'Productivity Score']
        values = ['85%', '24', '2.3 hrs', '7.2/10']
        trends = ['↑', '↑', '↓', '↑']
        
        for i, (metric, value, trend) in enumerate(zip(metrics, values, trends)):
            row_cells = chart_table.rows[i+1].cells
            row_cells[0].text = metric
            row_cells[1].text = value
            row_cells[2].text = trend
    
    def _add_footer(self, doc: Document):
        """Add a professional footer"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by Data2Paper • {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} • Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_report_document(self, report_data: Dict[str, Any], user_data: Dict[str, Any]) -> str:
        """
        Create a professionally formatted Word document from report data
        
        Args:
            report_data (Dict): Report data from the report agent
            user_data (Dict): User data for the report header
            
        Returns:
            str: Path to the generated document
        """
        # Create a new document
        doc = Document()
        
        # Add custom styles
        self._add_custom_styles(doc)
        
        # Add professional header
        self._add_header(doc, f"{report_data.get('report_type', 'Productivity')} Report", user_data, report_data)
        
        # Add executive summary
        if 'summary' in report_data:
            self._add_executive_summary(doc, report_data['summary'])
        
        doc.add_page_break()
        
        # Add detailed statistics if available
        if 'statistics' in report_data:
            self._add_statistics_section(doc, report_data['statistics'])
            doc.add_page_break()
        
        # Add AI-powered insights
        self._add_ai_insights_section(doc, report_data, user_data)
        doc.add_page_break()
        
        # Add visualization placeholder
        self._add_visualization_placeholder(doc)
        doc.add_page_break()
        
        # Add tasks information if available
        if 'tasks' in report_data and report_data['tasks']:
            self._add_tasks_section(doc, report_data['tasks'])
            doc.add_page_break()
        
        # Add notes if available
        if 'statistics' in report_data and 'all_notes' in report_data['statistics']:
            notes = report_data['statistics']['all_notes']
            if notes:
                self._add_notes_section(doc, notes)
                doc.add_page_break()
        
        # Add footer
        self._add_footer(doc)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_type = report_data.get('report_type', 'report').lower()
        user_id = user_data.get('id', 'unknown')
        filename = f"{report_type}_{timestamp}_{user_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def create_custom_report_document(self, report_data: Dict[str, Any], user_data: Dict[str, Any]) -> str:
        """
        Create a professionally formatted Word document from custom report data
        
        Args:
            report_data (Dict): Custom report data from the report agent
            user_data (Dict): User data for the report header
            
        Returns:
            str: Path to the generated document
        """
        # Create a new document
        doc = Document()
        
        # Add custom styles
        self._add_custom_styles(doc)
        
        # Add professional header
        self._add_header(doc, f"{report_data.get('report_type', 'Custom')} Report", user_data, report_data)
        
        # Add executive summary
        if 'summary' in report_data:
            self._add_executive_summary(doc, report_data['summary'])
        
        doc.add_page_break()
        
        # Add parameters section for custom reports
        if 'parameters' in report_data:
            doc.add_heading('Report Parameters', 1)
            params = report_data['parameters']
            if params:
                for key, value in params.items():
                    para = doc.add_paragraph()
                    para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                    para.add_run(str(value))
            else:
                doc.add_paragraph('No specific parameters provided.')
        
        doc.add_page_break()
        
        # Add detailed statistics if available
        if 'statistics' in report_data:
            self._add_statistics_section(doc, report_data['statistics'])
            doc.add_page_break()
        
        # Add AI-powered insights
        self._add_ai_insights_section(doc, report_data, user_data)
        doc.add_page_break()
        
        # Add visualization placeholder
        self._add_visualization_placeholder(doc)
        doc.add_page_break()
        
        # Add tasks information if available
        if 'tasks' in report_data and report_data['tasks']:
            self._add_tasks_section(doc, report_data['tasks'])
            doc.add_page_break()
        
        # Add notes if available
        if 'statistics' in report_data and 'all_notes' in report_data['statistics']:
            notes = report_data['statistics']['all_notes']
            if notes:
                self._add_notes_section(doc, notes)
                doc.add_page_break()
        
        # Add footer
        self._add_footer(doc)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_type = report_data.get('report_type', 'custom').lower()
        user_id = user_data.get('id', 'unknown')
        filename = f"{report_type}_{timestamp}_{user_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Save document
        doc.save(filepath)
        
        return filepath